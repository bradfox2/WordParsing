"""Fetch documents from SWMS Media"""

import inspect
import os
from pathlib import Path, PureWindowsPath
import warnings

import typed_ast
from docx import Document as docx_doc
from docx import styles

from dotenv import load_dotenv

from wip.convert import save_as_docx

load_dotenv(verbose=True, override=True)
DEBUG = os.getenv('DEBUG')
SWMS_MEDIA_FILER_PATH = os.getenv('SWMS_MEDIA_FILER_PATH')
SWMS_MEDIA_USER_NAME = os.getenv('SWMS_MEDIA_USER_NAME')
SWMS_MEDIA_USER_PASSWORD = os.getenv('SWMS_MEDIA_USER_PASSWORD')


class SWMSDocument(object):
    @classmethod
    @property
    def swms_media_path(cls):
        if not hasattr(cls, '_swms_media_path'):
            warnings.warn("No explicit SWMS_MEDIA_FILER_PATH set, using environment variable SWMS_MEDIA_FILER_PATH")
            cls.mount_swms_media_path(SWMS_MEDIA_FILER_PATH)
        return cls._swms_media_path

    @classmethod
    @swms_media_path.setter
    def swms_media_path(cls, val):
        cls._swms_media_path = val

    @classmethod
    def all_from_wm_id(cls, wm_id, dbcon):
        """ Returns a list of SWMSDocument objects related to the provided WM_ID """
        #TODO: get media ids from swms
        media_ids = []
        return [cls.from_swms_media_id(cls, x, dbcon) for x in media_ids]

    @classmethod
    def from_swms_media_id(cls, media_id, dbcon):
        #TODO: use dbcon to query SWMS, get path
        path = None
        return cls.from_file_path(path)

    @classmethod
    def from_file_path(cls, file_path):
        return SWMSDocument(Document(file_path))

    def __init__(self, document):
        self.document = document

    @classmethod
    def _check_media_path_mounted(cls, media_path):
        bad_strings = ['$','&']
        if any(bs in media_path for bs in bad_strings):
            raise ValueError('Bad string in swms media path.')
        cls._swms_media_path = Path(SWMS_MEDIA_FILER_PATH)
        if os.path.isdir(cls._swms_media_path) and not DEBUG:
            print("Media storage already connected.")
            return True
        return False

    @staticmethod
    def _windows_mount_command(user, password, path):
        return f"net use /user:{user} {path} {password}"

    @classmethod
    def _prep_linux_mount_command(user, password, path):
        media_path = f'/mnt{path[1:]}' if path[:2] =='//' else f'/mnt{path}'
        if not os.path.isdir(f'/tmp{media_path}'):
            os.system(f'mkdir /tmp{media_path}')
        return f'sudo mount -t cifs {path} {media_path} -o user={user},password={password},domain=nixcraft'


    @classmethod
    def mount_swms_media_path(cls, SWMS_MEDIA_FILER_PATH):
        if not cls._check_media_path_mounted(SWMS_MEDIA_FILER_PATH):
            print("Connecting to backup storage.")
            if os.name == 'nt':
                mount_command = _windows_mount_command(SWMS_MEDIA_USER_NAME, SWMS_MEDIA_USER_PASSWORD, SWMS_MEDIA_FILER_PATH)
            else:
                #untested
                mount_command = _prep_linux_mount_command(SWMS_MEDIA_USER_NAME, SWMS_MEDIA_USER_PASSWORD, SWMS_MEDIA_FILER_PATH)
            if DEBUG:
                print(mount_command)
            else:
                os.system(mount_command)
            media_dir = os.path.isdir(cls._swms_media_path)
            if media_dir:
                print("Connection success.")
            else:
                raise Exception("Failed to mount media directory.")

mount = SWMSDocument.mount_swms_media_path(SWMS_MEDIA_FILER_PATH)

class WorkInstruction(SWMSDocument):
    def __init__(self, wm_id, wt_id, rev, type, file_path, SWMS_MEDIA_ID):
        SWMSDocument.__init__(file_path)
        self.file_path = file_path
        self.wm_id
        self.wt_it
        self.rev
        self.type
        self.steps = []     

swms_media = os.path.isdir(SWMS_MEDIA_FILER_PATH)



class Document(object):
    def __init__(self, file):
        self.docx_obj = docx_doc(file)

class SWMSDocument(Document):
    def __init__(self, SWMS_SERVER_PATH):
        self.smid = SWMS_SERVER_PATH
        assert os.name == 'nt'      
       
    @staticmethod
    def fetch_documents(docs, conn):
        raise NotImplementedError

class WorkInstruction(SWMSDocument):
    def __init__(self, SWMS_MEDIA_ID):
        pass


if __name__ == '__main__':
    pass
