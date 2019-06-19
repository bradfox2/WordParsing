"""Fetch document from SWMS Media and return a path to the file"""
import importlib
import inspect
import os
import warnings
from os.path import exists
from pathlib import Path, PureWindowsPath

import typed_ast
from docx import Document as docx_doc
from docx import styles
from dotenv import load_dotenv

from sqlalchemy import create_engine, Column
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import Session
from sqlalchemy.dialects.sqlite import (
    BLOB, BOOLEAN, CHAR, DATE, DATETIME, DECIMAL, FLOAT, INTEGER, JSON,
    NUMERIC, SMALLINT, TEXT, TIME, TIMESTAMP, VARCHAR)

from wordparsing.utils import count_files, count_files_fast

from dotenv import load_dotenv

load_dotenv(verbose=True, override=True)

NIMS_CONNECTION_STRING = os.getenv('NIMS_CONNECTION_STRING')
engine = create_engine(NIMS_CONNECTION_STRING)
Base = declarative_base()

class NimsMediaDetail(Base):
    __tablename__ = 'media_details'
    db_id = Column('db_id', INTEGER, primary_key=True)
    server_path = Column('server_path', VARCHAR(length=2000))

class Document(object):
    # TODO: need a factory here to handle different types of documents.
    def __init__(self, fle):
        self.file = fle

class RemoteMount(object):
    '''Mount a remote file system'''

    def __init__(self, user, password, mount_path, remote_path):
        self.mount_path = mount_path
        self.filer_path = remote_path
        self.user = user
        self.password = password
        self._mount_path = None if mount_path is  None else Path(mount_path)
        self._remote_path = None if remote_path is None else Path(remote_path)
        self._mounted_path = None
        self.mount_remote()
    
    @property
    def mounted_path(self):
        return self._mounted_path.as_posix()

    def mount_remote(self):
        if os.name == 'nt':
            if not os.path.exists(self._remote_path):
                print("Mounting...")
                mount_command = self._prep_windows_mount_command(self.user, self.password, self._remote_path)
                self._mounted_path = self._remote_path
                print('Mounted.')
            else:
                warnings.warn("Remote already mounted.")
                self._mounted_path = self._remote_path
                return
        else:
            #untested
            mount_command = self._prep_linux_mount_command(self.user, self.password, self._mount_path, self._remote_path)
            self._mounted_path = self._mount_path.joinpath(self._remote_path) 
        
        os.system(mount_command)
        if self._mounted_path.exists():
            print("Connection success.")
        else:
            raise Exception("Failed to mount media directory.")
        return self._mounted_path
    
    def unmount_remote(self, mount_path):
        '''Unmount the remote filer'''
        raise NotImplementedError
    
    @property
    def not_allowed_strings(self):
        return ['$','&']

    def _sanitize_paths(self):
        if any(bs in self.filer_path for bs in self.not_allowed_strings):
            raise ValueError('Bad string in filer path.')

    def _check_filer_path_mounted(self):
        if self._mounted_path.exists():
            return True
        return False
    
    def _prep_linux_mount_command(self, user, password, mount_path, remote_path):
        #not tested
        return f'sudo mount -t cifs {mount_path} {remote_path} -o user={user},password={password},domain=nixcraft'

    def _prep_windows_mount_command(self, user, password, remote_path):
        # handle str since Path with UNC drives does not appear to want to work properly, 
        # and windows does not want to allow a connection to a domain under a different username.
        # mounting just the root drive appears to work though. 
        remote_path = [i for i in str(SWMS_MEDIA_FILER_PATH).split("\\") if i is not ""][0]
        return f"net use /user:{user} {remote_path.drive} {password}"

class SWMSDocument(Document):

    @classmethod
    def all_from_wm_id(cls, wm_id, engine, remote_mount):
        """ Returns a list of SWMSDocument objects related to the provided WM_ID """
        #TODO: map relationships into sql alchemy classes
        with open('queries\media_details_from_wmid.sql') as f:
            stmt = (''.join(f.readlines()))
        
        with engine.connect() as con:
            data = {"wmid":wm_id}
            rs = con.execute(stmt, **data)
            media_ids = list(rs)

        return [cls.from_filer_file_name(remote_mount, file_name, media_dbid) for wmwt_dbid, file_name, media_dbid in media_ids]

    @classmethod
    def all_from_raw_query(cls, query_file_path, filter_dict, engine, remote_mount):
        """ Returns a list of SWMSDocument objects related to the provided WM_ID.
            Query needs to return a WMWT dbid, a file name, and a media dbid in that order.
            Filter dict should have keys of the parameterized filters, and values of the passed parameters.
        """
        #TODO: Unpack the iterator of unknown parameters instead of requiring hard coded.
        #TODO: map relationships into sql alchemy classes

        with open(query_file_path) as f:
            stmt = (''.join(f.readlines()))
        with engine.connect() as con:
            rs = con.execute(stmt, **filter_dict)
            media_ids = list(rs)

        return [cls.from_filer_file_name(remote_mount, file_name, media_dbid) for wmwt_dbid, file_name, media_dbid in media_ids]

    @classmethod
    def from_swms_media_dbid(cls, media_dbid, session, remote_mount):
        media = session.query(NimsMediaDetail).filter_by(db_id = media_dbid).first()
        full_path = remote_mount.mounted_path + media.server_path
        return cls.from_absolute_file_path(full_path, media_dbid)

    @classmethod
    def from_absolute_file_path(cls, file_path, media_dbid):
        '''fetch file from absolute path'''
        full_path = Path(file_path)
        return(SWMSDocument(Document(full_path), media_dbid))

    @classmethod
    def from_filer_file_name(cls, remote_mount, file_name, media_dbid):
        '''fetch file from remote mount and file name'''
        full_path = remote_mount.mounted_path + file_name
        return(SWMSDocument(Document(full_path), media_dbid))
    
    def __init__(self, document, media_dbid):
        self.document = document
        self.path = document.file
        self.media_dbid = media_dbid

class WorkInstruction(SWMSDocument):
    def __init__(self, wm_id, wt_id, rev, typ, file_path, SWMS_MEDIA_ID):
        #TODO: Refactor SWMSDocument into traditional factory pattern.
        super().__init__(Document(file_path))
        self.file_path = file_path 
        self.wm_id = wm_id
        self.wt_it = wt_id
        self.rev = rev
        self.typ = typ
        self.steps = []     
    
class WorkInstructionStep(SWMSDocument):
    #TODO: Refactor fetching classes and types vs data_classes types, and doc_types as they are stored in the db.
    pass

if __name__ == '__main__':

    load_dotenv(verbose=True, override=True)
    DEBUG = os.getenv('DEBUG')
    SWMS_MEDIA_FILER_PATH = Path(os.getenv('SWMS_MEDIA_FILER_PATH'))
    SWMS_MEDIA_USER_NAME = os.getenv('SWMS_MEDIA_USER_NAME')
    SWMS_MEDIA_USER_PASSWORD = os.getenv('SWMS_MEDIA_USER_PASSWORD')

    #load NIMS connection information
    session = Session(engine)

    smws_filer_mount = RemoteMount(SWMS_MEDIA_USER_NAME,
                                SWMS_MEDIA_USER_PASSWORD,
                                None,
                                SWMS_MEDIA_FILER_PATH)

    a = SWMSDocument.from_swms_media_dbid(9322464, session, smws_filer_mount)

    #return: 
    b = SWMSDocument.all_from_wm_id(5144071, engine, smws_filer_mount)

